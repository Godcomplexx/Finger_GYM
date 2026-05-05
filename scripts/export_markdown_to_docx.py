from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "docx"

FONT_NAME = "Times New Roman"
BASE_FONT_SIZE = Pt(14)
SMALL_FONT_SIZE = Pt(12)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Лист ")
    run.font.name = FONT_NAME
    run.font.size = SMALL_FONT_SIZE

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def apply_document_format(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    for style_name in ("Normal", "Body Text"):
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = BASE_FONT_SIZE
        paragraph_format = style.paragraph_format
        paragraph_format.first_line_indent = Cm(1.25)
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = Pt(0)

    for style_name, size, bold in (
        ("Heading 1", 16, True),
        ("Heading 2", 15, True),
        ("Heading 3", 14, True),
    ):
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def format_paragraph(paragraph, *, indent: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = BASE_FONT_SIZE


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            run = paragraph.add_run(part)
        run.font.size = BASE_FONT_SIZE


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        return rows[0], rows[2:]
    return rows[0], rows[1:]


def add_table(document: Document, lines: list[str]) -> None:
    header, rows = parse_table(lines)
    table = document.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF7")
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = SMALL_FONT_SIZE

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(header)]):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, text)
            for run in p.runs:
                run.font.size = SMALL_FONT_SIZE
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0

    document.add_paragraph()


def add_code_block(document: Document, lines: list[str]) -> None:
    text = "\n".join(lines).rstrip()
    if not text:
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = SMALL_FONT_SIZE


def add_list_item(document: Document, text: str, numbered: bool = False) -> None:
    prefix = "" if numbered else "– "
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, prefix + text)


def convert_markdown(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    apply_document_format(document)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if is_table_line(line):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = document.add_heading(level=level)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.clear()
            run = p.add_run(text)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(16 if level == 1 else 15 if level == 2 else 14)
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            add_list_item(document, bullet_match.group(1), numbered=False)
            i += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_match:
            add_list_item(document, numbered_match.group(1), numbered=True)
            i += 1
            continue

        p = document.add_paragraph()
        add_inline_runs(p, stripped)
        format_paragraph(p)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def output_name(markdown_path: Path) -> str:
    relative = markdown_path.relative_to(ROOT)
    if relative.parent == Path("."):
        return f"ROOT_{markdown_path.stem}.docx"
    return f"{markdown_path.stem}.docx"


def main() -> None:
    markdown_files = [ROOT / "README.md", ROOT / "PLAN_DEVELOPMENT.md"]
    markdown_files.extend(sorted((ROOT / "docs").glob("*.md")))

    for markdown_path in markdown_files:
        output_path = OUTPUT_DIR / output_name(markdown_path)
        convert_markdown(markdown_path, output_path)
        print(f"{markdown_path.relative_to(ROOT)} -> {output_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
